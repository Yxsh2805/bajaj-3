import os
import time
import logging
import hashlib
import tempfile
import email
from datetime import datetime
from typing import List, Optional, Dict, Any
from urllib.parse import urlparse
from contextlib import asynccontextmanager
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document as DocxDocument
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

from fastapi import FastAPI, HTTPException, Header, Depends
from pydantic import BaseModel

# RAG imports
from langchain_together import ChatTogether, TogetherEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Expected Bearer token for authentication
EXPECTED_TOKEN = "5aa05ad358e859e92978582cde20423149f28beb49da7a2bbb487afa8fce1be8"

# ----- Request/Response Models -----
class QuestionRequest(BaseModel):
    documents: str
    questions: List[str]

class AnswerResponse(BaseModel):
    answers: List[str]

# ----- Simple Vector Store -----
class SimpleVectorStore:
    def __init__(self, embeddings_model):
        self.embeddings_model = embeddings_model
        self.documents = []
        self.embeddings = []
    
    def add_documents(self, docs: List[Document]):
        """Add documents to the vector store"""
        for doc in docs:
            self.documents.append(doc)
            # Get embedding for the document
            embedding = self.embeddings_model.embed_query(doc.page_content)
            self.embeddings.append(embedding)
    
    def similarity_search(self, query: str, k: int = 5) -> List[Document]:
        """Search for similar documents"""
        if not self.documents:
            return []
        
        # Get query embedding
        query_embedding = self.embeddings_model.embed_query(query)
        
        # Calculate similarities
        similarities = cosine_similarity([query_embedding], self.embeddings)[0]
        
        # Get top k indices
        top_indices = np.argsort(similarities)[::-1][:k]
        
        # Return top documents
        return [self.documents[i] for i in top_indices]

# ----- Multi-format Document Loader -----
def get_file_extension(url: str) -> str:
    """Extract file extension from URL"""
    parsed_url = urlparse(url)
    path = parsed_url.path.lower()
    if path.endswith('.pdf'):
        return 'pdf'
    elif path.endswith('.docx'):
        return 'docx'
    elif path.endswith('.eml'):
        return 'eml'
    else:
        return 'html'  # Default to HTML

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            
            text = ""
            with open(temp_file.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            os.unlink(temp_file.name)
            return text.strip()
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            
            doc = DocxDocument(temp_file.name)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            os.unlink(temp_file.name)
            return text.strip()
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise

def extract_text_from_eml(file_content: bytes) -> str:
    """Extract text from EML file"""
    try:
        msg = email.message_from_bytes(file_content)
        text = ""
        
        # Extract subject
        subject = msg.get('Subject', '')
        if subject:
            text += f"Subject: {subject}\n\n"
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                elif part.get_content_type() == "text/html":
                    html_content = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    text += soup.get_text()
        else:
            if msg.get_content_type() == "text/plain":
                text += msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            elif msg.get_content_type() == "text/html":
                html_content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                soup = BeautifulSoup(html_content, 'html.parser')
                text += soup.get_text()
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting EML text: {e}")
        raise

def extract_text_from_html(file_content: bytes) -> str:
    """Extract text from HTML content"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text content
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        text = ' '.join(text.split())
        
        return text
    except Exception as e:
        logger.error(f"Error extracting HTML text: {e}")
        raise

def load_document_content(url: str) -> List[Document]:
    """Load and process document of various formats (PDF, DOCX, EML, HTML)"""
    try:
        # Download the file
        response = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
        response.raise_for_status()
        
        file_content = response.content
        file_extension = get_file_extension(url)
        
        logger.info(f"Processing {file_extension.upper()} file from {url}")
        
        # Extract text based on file type
        if file_extension == 'pdf':
            text = extract_text_from_pdf(file_content)
        elif file_extension == 'docx':
            text = extract_text_from_docx(file_content)
        elif file_extension == 'eml':
            text = extract_text_from_eml(file_content)
        else:  # HTML or unknown
            text = extract_text_from_html(file_content)
        
        if not text.strip():
            raise ValueError(f"No text content extracted from {file_extension.upper()} file")
        
        logger.info(f"Successfully extracted {len(text)} characters from {file_extension.upper()} file")
        
        return [Document(page_content=text, metadata={"source": url, "file_type": file_extension})]
        
    except Exception as e:
        logger.error(f"Failed to load document from {url}: {e}")
        raise

# ----- Railway-Optimized RAG Engine -----
class RailwayRAGEngine:
    def __init__(self):
        self.chat_model = None
        self.embeddings = None
        self.text_splitter = None
        self.policy_prompt = None
        self.initialized = False
        
        # Minimal caching for Railway
        self.document_cache: Dict[str, Any] = {}
        self.max_cache_size = 2
    
    def _get_url_hash(self, url: str) -> str:
        """Generate hash for URL for caching purposes"""
        return hashlib.md5(url.encode()).hexdigest()[:12]
    
    def initialize(self):
        """Initialize RAG components for Railway environment"""
        if self.initialized:
            return
            
        logger.info("Initializing Railway RAG engine...")
        
        try:
            # Set environment variables
            os.environ["TOGETHER_API_KEY"] = os.getenv("TOGETHER_API_KEY", "deb14836869b48e01e1853f49381b9eb7885e231ead3bc4f6bbb4a5fc4570b78")
            os.environ["LANGCHAIN_TRACING_V2"] = "true"
            os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGCHAIN_API_KEY", "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")
            os.environ["LANGCHAIN_PROJECT"] = "railway-rag-deployment"

            # Initialize LLM and embeddings
            self.chat_model = ChatTogether(
                model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
                temperature=0,
                max_tokens=3000
            )
            self.embeddings = TogetherEmbeddings(model="BAAI/bge-base-en-v1.5")

            # Optimized text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=600,
                chunk_overlap=80,
                separators=["\n\n", "\n", ". ", " ", ""]
            )

            # Prompt template
            self.policy_prompt = ChatPromptTemplate([
                ("system", """You are an expert document assistant. Answer questions concisely and accurately based on the provided document content.

CRITICAL FORMAT: Input questions are separated by " | ". Output answers MUST be separated by " | " in the same order.

Guidelines:
- Direct, concise answers
- Lead with key information
- Cite specific document content when available
- If unsure, state limitations clearly
- Maintain exact order and use " | " separator between answers"""),
                ("human", """Questions: {query}
Context: {context}"""),
            ])

            self.initialized = True
            logger.info("Railway RAG engine initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize RAG engine: {str(e)}")
            raise

    def _load_and_process_document(self, url: str) -> tuple:
        """Load and process document with multi-format support"""
        if url in self.document_cache:
            logger.info(f"Using cached document for {url}")
            return self.document_cache[url]
        
        logger.info(f"Loading document: {url}")
        start_time = time.time()
        
        try:
            # Use multi-format loader
            docs = load_document_content(url)
            chunks = self.text_splitter.split_documents(docs)
            
            load_time = time.time() - start_time
            logger.info(f"Document loaded and chunked in {load_time:.2f}s ({len(chunks)} chunks)")
            
            # Cache management
            if len(self.document_cache) >= self.max_cache_size:
                oldest_key = next(iter(self.document_cache))
                del self.document_cache[oldest_key]
            
            self.document_cache[url] = (docs, chunks)
            return docs, chunks
            
        except Exception as e:
            logger.error(f"Failed to load document {url}: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to load document: {str(e)}")

    def _create_vectorstore(self, chunks: List) -> SimpleVectorStore:
        """Create simple vector store"""
        logger.info("Creating simple vector store")
        start_time = time.time()
        
        try:
            vectorstore = SimpleVectorStore(self.embeddings)
            vectorstore.add_documents(chunks)
            
            creation_time = time.time() - start_time
            logger.info(f"Vector store created in {creation_time:.2f}s")
            
            return vectorstore
            
        except Exception as e:
            logger.error(f"Vector store creation error: {e}")
            raise HTTPException(status_code=500, detail="Failed to create vector store")

    def process_document_questions(self, url: str, questions: List[str]) -> List[str]:
        """Process document and answer questions"""
        if not self.initialized:
            raise RuntimeError("RAG engine not initialized")
        
        total_start_time = time.time()
        
        try:
            # Load and process document
            docs, chunks = self._load_and_process_document(url)
            
            # Create vector store
            vectorstore = self._create_vectorstore(chunks)
            
            # Process questions in batch
            logger.info(f"Processing {len(questions)} questions in batch...")
            batch_query = " | ".join(questions)
            
            # Get relevant context
            relevant_docs = vectorstore.similarity_search(batch_query, k=5)
            context = " ".join([doc.page_content for doc in relevant_docs])[:3000]
            
            # Create prompt
            prompt = self.policy_prompt.format(query=batch_query, context=context)
            
            # Get response from LLM
            query_start_time = time.time()
            batch_result = self.chat_model.invoke(prompt)
            query_time = time.time() - query_start_time
            
            logger.info(f"LLM query completed in {query_time:.2f}s")
            
            # Parse results
            if hasattr(batch_result, 'content'):
                batch_result = batch_result.content
            
            answers = [answer.strip() for answer in str(batch_result).split(" | ")]
            
            # Validate answer count
            if len(answers) != len(questions):
                logger.warning(f"Answer count mismatch: {len(questions)} questions, {len(answers)} answers")
                while len(answers) < len(questions):
                    answers.append("Unable to generate answer for this question.")
                answers = answers[:len(questions)]
            
            total_time = time.time() - total_start_time
            logger.info(f"Total processing time: {total_time:.2f}s")
            
            return answers

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise

# Global RAG engine instance
rag_engine = RailwayRAGEngine()

# ----- Token Verifier -----
def verify_token(authorization: Optional[str] = Header(None)):
    if authorization is None or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Authorization header missing or invalid format")
    
    token = authorization.split("Bearer ")[-1]
    if token != EXPECTED_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid Bearer token")

# ----- Lifespan Context Manager -----
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Lifespan context manager for startup and shutdown"""
    # Startup
    try:
        rag_engine.initialize()
        logger.info("Railway application startup completed successfully")
    except Exception as e:
        logger.error(f"Startup error: {str(e)}")
        logger.info("Continuing with limited functionality...")
    
    yield
    
    # Shutdown
    try:
        # Clean up resources if needed
        rag_engine.document_cache.clear()
        logger.info("Application shutdown completed")
    except Exception as e:
        logger.error(f"Shutdown error: {str(e)}")

# ----- FastAPI App -----
app = FastAPI(title="Railway RAG API", version="1.0.0", lifespan=lifespan)

@app.post("/hackrx/run", response_model=AnswerResponse)
async def ask_questions(
    request: QuestionRequest,
    authorization: str = Depends(verify_token)
):
    """Main endpoint for processing documents and answering questions"""
    try:
        logger.info(f"Received request with {len(request.questions)} questions")

        # Validation
        if not request.documents.startswith(('http://', 'https://')):
            raise HTTPException(status_code=400, detail="Invalid document URL")
        if not request.questions:
            raise HTTPException(status_code=400, detail="Questions list is empty")

        # Process questions
        answers = rag_engine.process_document_questions(request.documents, request.questions)

        logger.info(f"Successfully processed {len(answers)} answers")
        return {"answers": answers}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Internal processing error: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "rag_initialized": rag_engine.initialized,
        "cached_documents": len(rag_engine.document_cache),
        "timestamp": datetime.now().isoformat()
    }

# Railway startup
if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))  # Railway sets PORT automatically
    uvicorn.run(app, host="0.0.0.0", port=port)
